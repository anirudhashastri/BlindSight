from docx import Document
from groq import Groq
import os
from dotenv import load_dotenv
from pynput import keyboard
import threading
import sys
from log_config import setup_logger
from TTS import speak, set_speed  
import sys
import os
from datetime import datetime

# Create logger for the specific module
logger = setup_logger("documentreader")
logger.info(f"Logging initialized for {__name__}")

# Load environment variables
load_dotenv()

# Initialize Groq client with API key from environment variables
groq_api_key = os.getenv("GROQ_API_KEY")
if not groq_api_key:
    logger.error("GROQ_API_KEY not found in environment variables.")
    sys.exit(1)

try:
    client = Groq(api_key=groq_api_key)
    logger.info("Groq client initialized successfully.")
except Exception as e:
    logger.error(f"Failed to initialize Groq client: {e}")
    sys.exit(1)

def read_docx(filename):
    """
    Reads the content of a .docx or .txt file.
    
    Args:
        filename (str): Path to the file.
    
    Returns:
        str or Document: The content of the file. Returns a Document object for .docx files
                        and a string for .txt files.
    """    
    try:
        if filename.lower().endswith(".docx"):
            document = Document(filename)
            logger.info(f"Successfully read .docx file: {filename}")
            return document
        elif filename.lower().endswith(".txt"):
            with open(filename, "r", encoding='utf-8') as file:
                content = file.read()
                logger.info(f"Successfully read .txt file: {filename}")
                return content
        else:
            logger.warning(f"Unsupported file format for file: {filename}")
            return None
    except FileNotFoundError:
        logger.error(f"File not found: {filename}")
        return None
    except Exception as e:
        logger.error(f"Error reading file {filename}: {e}")
        return None

def write_docx(file_path, content):
    """
    Writes content to a .docx or .txt file.
    
    Args:
        file_path (str): Path to the file to be written.
        content (str or Document): The content to write. If writing to a .docx file, 
                                   'content' should be a Document object.
    
    Raises:
        ValueError: If the file extension is unsupported.
    """
    try:
        if file_path.lower().endswith(".docx") and isinstance(content, type(Document)):
            content.save(file_path)
            logger.info(f"Successfully wrote to .docx file: {file_path}")
        elif file_path.lower().endswith(".txt") and isinstance(content, str):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(content)
                logger.info(f"Successfully wrote to .txt file: {file_path}")
        else:
            logger.error("Unsupported file format or content type.")
            raise ValueError("Unsupported file format or content type.")
    except Exception as e:
        logger.error(f"Error writing to file {file_path}: {e}")
        raise

def doc_operations(document, command):
    """
    Processes the document content with the Groq API based on the specified command.
    
    Args:
        document (str or Document): The content of the document.
        command (str): The operation to perform on the document.
    
    Returns:
        str: The updated content of the document.
    """

    doc_content = document
    if hasattr(document, 'paragraphs'):
        doc_content = '\n'.join(para.text for para in document.paragraphs)
    elif not isinstance(document, str):
        doc_content = str(document)
    
    prompt = [
        {
            "role": "system",
            "content": """
            Given a document and a specified operation, apply the operation to the document
            and return only the updated text. Do not include any explanations or extra information.
            If it is only a blank audio for command. Just return the original document.
            Do the same for read document command.
            """
        },
        {
            "role": "user",
            "content": f"{doc_content}\n\nOperation:\n{command}"}
    ]

    try:
        response = client.chat.completions.create(
            messages=prompt,
            model="llama3-70b-8192"
        )
        
        result = response.choices[0].message.content.strip()
        logger.info("Document operation completed successfully.")
        return result
    
    except Exception as e:
        logger.error(f"Error during API call in doc_operations: {e}")
        return None

def doc_reading(command):
    """
    Extracts the filename with its extension from the given command using Groq API.
    
    Args:
        command (str): The user command containing the filename.
    
    Returns:
        str: The extracted filename with extension.
    """    
    prompt = [
        {
            "role": "system",
            "content": """
                Extract the filename with its extension from the given command. If a full file path 
                is provided, return the entire file path including the extension. Respond with only 
                the result, no additional text or explanation.
                
                Eg:
                Command: Open my text file food.
                Response: food.txt
                
                Command: open file sunrize dot txt.
                Response: sunrize.txt
                
                Command: Open my document money.
                Response: money.docx
            """
        },
        {
            "role": "user",
            "content": command
        }
    ]

    try:
        response = client.chat.completions.create(
            messages=prompt,
            model="llama3-70b-8192"
        )
        
        result = response.choices[0].message.content.strip()
        logger.info(f"Extracted filename: {result}")
        return result
    
    except Exception as e:
        logger.error(f"Error during API call in doc_reading: {e}")
        return None

def doc_main(command, speech_recog):
    """
    Main function to handle document editing based on user commands.
    
    Args:
        command (str): The initial command to open a document.
        speech_recog (STT): An instance of the STT (Speech-to-Text) class.
    """
    # Set slightly faster speed for system messages
    speak("You are now in document editor mode. Say Exit at the end to close the document after all your changes.", speed=1.2)
    logger.info("Entered document editor mode.")
    
    filename = doc_reading(command)
    if not filename:
        speak("Could not extract the filename from your command.")
        logger.error("Filename extraction failed.")
        return
    
    logger.info(f"Filename extracted: {filename}")
    speak(f"Opening the document {filename}.")
    
    while True:
        document_content = read_docx(filename)
        if document_content is None:
            speak("There is no content to read. What would you like to do?")
        else:
            speak("Document is now open. What would you like to do?", speed=1.2)
        
        # Record audio input from the user
        speech_recog.record_audio("audio_sample.wav", duration=7)
        command = speech_recog.process_audio_with_whisper("audio_sample.wav")
        logger.info(f"User command: {command}")
        
        if not command:
            speak("I did not catch that. Please repeat your command.")
            continue
        
        if "exit" in command.lower() or "close document" in command.lower():
            speak("Closing the document editor. Goodbye!", speed=1.2)
            logger.info("Exiting document editor mode.")
            break
        
        elif 'read' in command.lower():
            if isinstance(document_content, type(Document)):
                full_text = []
                for para in document_content.paragraphs:
                    full_text.append(para.text)
                text_to_read = '\n'.join(full_text)
            else:
                text_to_read = document_content
            # Use normal speed for document content
            speak(text_to_read, speed=1.0)
            logger.info("Read command executed.")
        
        else:
            updated_content = doc_operations(document_content, command)
            if updated_content is None:
                speak("Operation failed.")
                logger.error("Document operation failed.")
            elif any(keyword in command.lower() for keyword in ["summary", "summarize", "how many", "what is"]):
                # Use slightly faster speed for summaries
                speak(updated_content, speed=1.3)
                logger.info("Informational operation executed.")
            else:
                # It's an update operation; write back to the file
                if filename.lower().endswith(".docx"):
                    doc = Document()
                    for line in updated_content.split('\n'):
                        doc.add_paragraph(line)
                    write_docx(filename, doc)
                    speak(f"File '{filename}' has been updated.", speed=1.2)
                    logger.info(f"File '{filename}' has been updated.")
                elif filename.lower().endswith(".txt"):
                    write_docx(filename, updated_content)
                    speak(f"File '{filename}' has been updated.", speed=1.2)
                    logger.info(f"File '{filename}' has been updated.")
                else:
                    logger.error("Unsupported file extension.")
                    speak("Unsupported file format.")

def main():
    """
    Entry point of the script. Reads a document and initiates the TTS.
    """
    file_path = "test/text.txt"
    initial_command = "Change sun to moon"
    
    logger.info("Starting document reader script.")
    
    document_content = read_docx(file_path)
    if document_content is None:
        speak("Failed to read the document. Please check the file path and format.")
        logger.error("Initial document reading failed.")
        return
    
    if isinstance(document_content, type(Document)):
        full_text = []
        for para in document_content.paragraphs:
            full_text.append(para.text)
        text_to_speak = '\n'.join(full_text)
    else:
        text_to_speak = document_content
    
    speak(text_to_speak)
    logger.info("Initial document content spoken.")
    
    try:
        from STT import STT
        speech_recog = STT()
        logger.info("Speech-to-Text system initialized.")
    except ImportError as e:
        logger.error(f"Failed to import STT class: {e}")
        speak("Speech recognition system is not available.")
        return
    except Exception as e:
        logger.error(f"Error initializing STT system: {e}")
        speak("An error occurred while initializing the speech recognition system.")
        return
    
    doc_main(initial_command, speech_recog)
    logger.info("Document reader script finished.")

if __name__ == "__main__":
    file_path = "test/sunrise.txt"
    command = "Change sun to moon"
    document_content = read_docx(file_path)
    speak(document_content)