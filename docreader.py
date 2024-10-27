
import speech_recognition as sr
import pyttsx3
from docx import Document
import ollama
import os







# Initialize TTS engine
engine = pyttsx3.init()


# Function to speak text
def speak(text):
    engine.say(text)
    engine.runAndWait()

# Listen to voice commands
def listen(mic_index=None):
    recognizer = sr.Recognizer()
    mic = sr.Microphone(device_index=mic_index) if mic_index is not None else sr.Microphone()
    
    with mic as source:
        print("Listening for command...")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)
        
        try:
            command = recognizer.recognize_google(audio)
            print(f"You said: {command}")
            return command.lower()
        except sr.UnknownValueError:
            speak("Sorry, I didn't catch that.")
            return ""
        except sr.RequestError:
            speak("Sorry, there was a network error.")
            return ""




# Intent recognition using LLM
def recognize_intent_with_llm(command):
    prompt = (
        f"User command: \"{command}\"\n"
        "Identify the intent of the command and respond with one of these labels:\n"
        "add_text, delete_line, read_document, read_line, next_line, previous_line, edit_line, exit.\n"
        "Respond only with the intent label and nothing else."
    )
    # Send prompt to LLM
    stream = ollama.chat(
        model='llama3.1:8b',
        messages=[{'role': 'user', 'content': prompt}],
        stream=True
    )
    
    intent = ""
    for chunk in stream:
        intent += chunk['message']['content']
    
    return intent.strip().lower()



# Function to extract line number using LLM
def get_line_number_with_llm(command):
    prompt = (
        f"User input: \"{command}\"\n"
        "Identify and respond with only the numeric line number mentioned in the input. "
        "For example, if the input is 'read line three,' respond with '3'. No extra text, just the number."
    )
    # Send prompt to LLM
    stream = ollama.chat(
        model='llama3.1:8b',
        messages=[{'role': 'user', 'content': prompt}],
        stream=True
    )
    
    line_number = ""
    for chunk in stream:
        line_number += chunk['message']['content']
    
    return line_number.strip()


# Open and initialize document
def open_document(filename):
    if filename.endswith(".docx"):
        return Document(filename)
    elif filename.endswith(".txt"):
        with open(filename, "r+") as file:
            return file.read().splitlines()  # Treat each line as an entry in a list


# Save changes to the document
def save_document(content, filename):
    if filename.endswith(".docx"):
        content.save(filename)
    elif filename.endswith(".txt"):
        with open(filename, "w") as file:
            file.write("\n".join(content))  # Save content list back to file


# Add text to the document
def add_text(content, filename, new_text):
    if filename.endswith(".docx"):
        content.add_paragraph(new_text)
    elif filename.endswith(".txt"):
        content.append(new_text)
    speak("Text added successfully.")

# Delete text by line number
def delete_line(content, line_num, filename):
    try:
        if filename.endswith(".docx"):
            paragraph = content.paragraphs[line_num]
            paragraph.clear()
        elif filename.endswith(".txt"):
            content.pop(line_num)
        speak("Line deleted successfully.")
    except IndexError:
        speak("Invalid line number.")

# Use LLM to edit a specific line
def edit_line_with_llm(original_line, edit_instruction):
    prompt = (
        f"Original line: \"{original_line}\"\n"
        f"Edit request: {edit_instruction}\n"
        "Provide only the updated line with no extra text:"
    )
    # Send prompt to LLM
    stream = ollama.chat(
        model='llama3.1:8b',
        messages=[{'role': 'user', 'content': prompt}],
        stream=True
    )
    
    edited_line = ""
    for chunk in stream:
        edited_line += chunk['message']['content']

    return edited_line.strip()


# Read and navigate through the document
def read_document(content, filename, line_num=None):
    if filename.endswith(".docx"):
        text = content.paragraphs[line_num].text if line_num is not None else "\n".join([p.text for p in content.paragraphs])
    elif filename.endswith(".txt"):
        text = content[line_num] if line_num is not None else "\n".join(content)
    
    speak(text)



def edit_document(mic_index, file_path):
    # Hardcoded file path for testing
    filename = file_path

    if not os.path.exists(filename):
        speak("File not found. Please try again.")
        return

    # Open the document
    content = open_document(filename)
    current_line = 0

    speak("File opened successfully. You can say commands like add text, delete line, edit line, or read document.")
    
    while True:
        command = listen(mic_index)
        intent = recognize_intent_with_llm(command)
        
        if intent == "exit":
            save_document(content, filename)
            speak("Changes saved. Goodbye!")
            break
        
        elif intent == "add_text":
            speak("What text would you like to add?")
            new_text = listen(mic_index)
            add_text(content, filename, new_text)

        elif intent == "delete_line":
            speak("Which line number would you like to delete?")
            try:
                line_num = int(get_line_number_with_llm(listen(mic_index)))
                delete_line(content, line_num - 1, filename)
            except ValueError:
                speak("Invalid line number.")

        elif intent == "read_document":
            read_document(content, filename)

        elif intent == "read_line":
            speak("Which line number would you like to hear?")
            try:
                line_command = listen(mic_index)
                line_num = int(get_line_number_with_llm(line_command))
                read_document(content, filename, line_num - 1)
            except ValueError:
                speak("Invalid line number.")
        
        elif intent == "next_line":
            current_line += 1
            read_document(content, filename, current_line)

        elif intent == "previous_line":
            current_line = max(0, current_line - 1)
            read_document(content, filename, current_line)
        
        elif intent == "edit_line":
            speak("Which line number would you like to edit?")
            try:
                line_command = listen(mic_index)
                line_num = int(get_line_number_with_llm(line_command)) - 1
                
                original_line = content.paragraphs[line_num].text if filename.endswith(".docx") else content[line_num]
                
                speak("What changes would you like to make?")
                edit_instruction = listen(mic_index)
                
                # Get edited line from LLM
                edited_line = edit_line_with_llm(original_line, edit_instruction)
                
                # Update the document with the edited line
                if filename.endswith(".docx"):
                    content.paragraphs[line_num].text = edited_line
                elif filename.endswith(".txt"):
                    content[line_num] = edited_line
                
                speak("Line edited successfully.")
            except (ValueError, IndexError):
                speak("Invalid line number or edit command.")
        
        else:
            speak("Sorry, I didn't understand that command. Please try again.")

